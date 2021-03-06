from django.conf import settings
from docx import Document
import datetime
import re
from PIL import Image, ImageOps
import fitz  # PyMuPDF
import io
import os
import re
import openpyxl as xl


# os.altsep = '\\'


class Fcc:

    def fig_captxn_citatxn(self, doc_name):
        try:
            doc = Document(r'{}'.format(doc_name))

            para_list = []
            all_text = []
            idx = []
            fcc = []
            count = 0
            caption = []
            citation = []
            Error_log = []

            # print(type(doc.paragraphs))

            for para in doc.paragraphs:
                if len(para.text) < 1:  # remove space between the lines
                    continue
                para.text = para.text.replace("\u2022 ", "")
                para.text = para.text.replace("\u2022", "")
               # para.text = para.text.replace(" ", " ")
                para.text = para.text.replace("\u2002", " ")
                para.text = para.text.replace("\u2002 ", " ")
                para.text = para.text.replace("\u2003", " ")
                para.text = para.text.replace("\u2003 ", " ")
                #para.text = para.text.replace("\"", " ")
                all_text.append(para.text)
                # print("------")

                #print(para.text)
                # print(all_text)
                if re.findall('^Fig', para.text):
                    count += 1
                    idx.append(all_text.index(para.text))  # index of words starting with Fig
                    #print(para.text)
                elif re.findall('^FIGURE', para.text):
                    count += 1
                    idx.append(all_text.index(para.text))  # index of words starting with Fig
                    #para.text = para.text.replace("\u0020", "")
                    #print(para.text)
                elif re.findall('^Figure', para.text):
                    count += 1
                    idx.append(all_text.index(para.text))
                elif re.findall('^FIG', para.text):
                    count += 1
                    idx.append(all_text.index(para.text))
                # index of words starting with Fig
                    #print(para.text)

                    # print("----")
            # print("indexes of Fig")
            # print(idx)
            max1 = 0
            x = []
            for i in range(len(idx) - 1):  # if the index distance is less than 10
                if (idx[i + 1] - idx[i]) > 10:  # if any other word got captured like figureus from upper portion of file
                    x.append(i)
                elif max1 < (idx[i + 1] - idx[i]) and (idx[i + 1] - idx[i]) < 10:  # if the index distance is less than 10
                    max1 = (idx[i + 1] - idx[i])

            # print("new indexes")
            # print(x)
            # print(idx)
            print(max1,"max")
            # idx.pop(i)
            for i in (x):  # removing extra index for idx
                idx.pop(i)
            # print(idx)

            # print("XYZ done")
            try:
                for i in range(idx[0], idx[-1] + max1):  # is used to prevent last fig's (capion and citation)
                    fcc.append(all_text[i])
            except:
                print("Check this file manually  (No fig found)!")
            #print(fcc[0],"fcc")
            #print(fcc[9], "fcc")

            idx1 = []
            for text in fcc:
                if re.findall('^Fig', text):
                    idx1.append(fcc.index(text))
                elif re.findall('^FIGURE', text):
                    idx1.append(fcc.index(text))
                elif re.findall('^FIG', text):
                    idx1.append(fcc.index(text))

            # print(idx1)
            figure = []
            flag_ = []
            y = all_text[0].split(" ")

            for j in idx:
                if re.findall('^figure', all_text[j].lower()):
                    col=6
                elif re.findall('^fig.', all_text[j].lower()):
                    col=4
                elif re.findall('^fig', all_text[j].lower()):
                    col=3
                l=len(all_text[j])
                #print(all_text[j],"jjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjjj")
                if all_text[j][:col].lower()== "figure" or all_text[j][:col].lower()=="fig" or all_text[j][:col].lower()== "fig.":
                    #print("yes")
                    if all_text[j][col]!=" ":
                        t=(list(all_text[j]))
                        t.insert(col," ")
                        all_text[j]="".join(t)
                    col+=1
                    #for z in range(col,len(all_text[j][col:])):

                    ss=re.findall('[A-Za-z]', all_text[j][col:])
                    #print("Ture",ss)
                    if len(ss)>0:
                        dd=all_text[j].split(ss[0])
                        d=all_text[j][col:len(dd[0])]
                        d = d.replace(" ", "")
                        all_text[j]=all_text[j][:col]+d+" "+ss[0]+dd[1]
                        #print(all_text[j],"dddd")
                    else:
                        #print(all_text[j])
                        d = all_text[j][col:]
                        d = d.replace(" ", "")
                        all_text[j] = all_text[j][:col] + d
                        #print(all_text[j],"ssss ",d,j)




                #print(all_text[j],"------------")



            for j in idx:
                y1 = all_text[j].split(" ")
                #print((y1[2])=="")
                #print(len(y1[2]))
                #print(all_text[j].split(" "))
                if y1 != y:
                    try:
                        if len(y1[2]) >= 0:
                            flag_.append(0)
                            figure.append(y1[0] + " " + y1[1])
                            caption.append(all_text[j][len(y1[0] + " " + y1[1]):])
                        # print(caption,"fvf")
                        else:
                            flag_.append(1)
                            figure.append(all_text[j])
                            if not (re.findall(("From" or "from" or 'FROM' or '(From' or '(from' or '(FROM' or '(Copyright'),all_text[j + 1])):
                                caption.append(all_text[j + 1])
                           # print(caption, "fyhjvf")

                                # print(len(y1[2]),"------",len(y1))
                    except:
                        # print(len(y1))
                        flag_.append(1)
                        figure.append(all_text[j])
                        if not (re.findall(("From" or "from" or 'FROM' or '(From' or '(from' or '(FROM' or '(Copyright'),all_text[j + 1])):
                            caption.append(all_text[j + 1])
                #print(caption[-1])            #print(caption,"dscs")

            for i in range(0, len(idx1)):
                k = idx1[i]
                # print(k,'k')
                try:
                    z = idx1[i + 1]
                    #print(z,'z')

                except:
                    z = len(fcc)

                token = 0
                for j in range(k, z):
                    if re.findall(("From" or "from" or 'FROM' or '(From' or '(from' or '(FROM' or '(Copyright'),fcc[j]):
                        token = 1
                        c = j
                if token == 1:
                    if fcc[c][0] == "(":
                        fcc[c] = fcc[c][1:-1]
                    citation.append(fcc[c])
                    #print(citation[-1],"fff")

                else:
                    citation.append("")

        # return figure, caption, citation

        except Exception as e:

            Error_log.append(e)
        # print(caption[0],"cit")
        # print(caption[9],"cit")


        #print(len(figure), " F ", len(caption), ", C ", len(citation))
        for i in range(len(figure)):
             print(figure[i][:20], ",", caption[i][:20], ",", citation[i][:20])
        # print(list(zip(figure, caption, citation)))
        return list(zip(figure, caption, citation))

        # for i in range(len(figure)):
            # print(figure[i][:20], ",", caption[i][:20], ",", citation[i][:20])

    def imgRenaming(self, otherImgList, isbn):
        figurePath = os.path.dirname(os.path.abspath(otherImgList[0]))
        otherImgList = [os.path.basename(i) for i in otherImgList]
        # fig=list(img)
        # print(fig,"**********")
        otherImgList = [i for i in otherImgList if (i.split(".")[-1] != "xlsx" and i.split(".")[-1] != 'pptx')]
        k=[]
        for fig in otherImgList:
            if re.findall('fig[A-Za-z]', fig.lower()):

                os.remove(r"{}/{}".format(figurePath, fig))
                k.append(fig)

           # print(fig, "dddddddddddddddd",figurePath)
        for i in k:
            otherImgList.remove(i)

        for fig in otherImgList:
            #     print(fig)
            #
            #     print(fig)

            # fig="CH0001_Figurunn005__v1_Orig.tif"
            # print(fig)
            if re.search("Figurunn", fig):
                k = "u"
                x = 8
            elif re.search("tb", fig):
                k = "t"
                x = 2
            elif re.search("Fig", fig):
                k = "f"
                x = 3
            else:
                k = "co"
                x = 0
            fig1 = fig.split("_")
           # print(fig1,"Aaaaaaaaa",fig)
            #

            if (len(fig1[1]) - x) <= 3:
                newFig = k + fig1[0][-2:] + "-" + fig1[1][-3:] + "-" + isbn + "." + fig.split('.')[-1]
            elif "." in fig1[1]:
                newFig = k + fig1[0][-2:] + "-" + fig1[1][x:x + 3] + "-" + fig1[1][x + 4:] + "-" + isbn + "." + \
                         fig.split('.')[-1]
            else:
                newFig = k + fig1[0][-2:] + "-" + fig1[1][x:x + 3] + "-" + fig1[1][x + 3:] + "-" + isbn + "." + \
                         fig.split('.')[-1]


            newFig1 = newFig.split(".")
            if newFig1[-1] not in "eps" and newFig1[-1] not in "tiff":
                newFig1[-1] = "tiff"
                newFig2 = ".".join(newFig1)
                #print(fig, "**************", newFig2)
                os.rename(r"{}/{}".format(figurePath, fig), r"{}/{}".format(figurePath, newFig2))
            else:
                # print(fig,"************",newFig)
                os.rename(r"{}/{}".format(figurePath, fig), r"{}/{}".format(figurePath, newFig))

        # 1. rename the current image

        # 2. delete old

    def pdfProcess(self, pdfList, isbn):
        # 1. extract images
        figurePath = os.path.dirname(os.path.abspath(pdfList[0]))
        pdfList = [os.path.basename(i) for i in pdfList]
        # print(r'{}\{}'.format(filePath,pdfList))

        for figPDF in pdfList:
            # print(figPDF,"((((")

            k = ""
            if re.search("Figurunn", figPDF):
                k = "u"
                x = 8
            elif re.search("tb", figPDF):
                k = "t"
                x = 2
            elif re.search("Fig", figPDF):
                k = "f"
                x = 3
            else:
                k = "co"
                x = 0
            figPDF1 = figPDF.split("_")
            # print(fig1,"Aaaaaaaaa",fig)
            #
            #
            if (len(figPDF1[1]) - x) <= 3:
                # print("if",k + figPDF1[0][-2:] + "-" + figPDF1[1][-3:] + "-" + isbn + ".pdf")
                newFigPdf = k + figPDF1[0][-2:] + "-" + figPDF1[1][-3:] + "-" + isbn + ".pdf"
            elif "." in figPDF1[1]:
                newFigPdf = k + figPDF1[0][-2:] + "-" + figPDF1[1][x:x + 3] + "-" + figPDF1[1][x + 4:] + "-" + isbn + ".pdf"
            else:
                # print("else",k + figPDF1[0][-2:] + "-" + figPDF1[1][x:x + 3] + "-" + figPDF1[1][x + 4:] + "-" + isbn + ".pdf")
                newFigPdf = k + figPDF1[0][-2:] + "-" + figPDF1[1][x:x + 3] + "-" + figPDF1[1][x + 3:] + "-" + isbn + ".pdf"
            # print(figPDF1,"((((((((",figPDF)
            # newFigPdf = k + figPDF1[0][-2:] + "-" + figPDF1[1][-2:] + "-" + isbn + ".pdf"
            # print(filePath,figPDF,"*********")
            # print(filePath, newFigPdf, "*********")
            os.rename(r"{}/{}".format(figurePath, figPDF), r"{}/{}".format(figurePath, newFigPdf))
            # os.remove(r"{}\{}".format(filePath, newFigPdf))
            newFigPdf1 = newFigPdf.split(".")
            # print(newFigPdf,"****")

            pdf_file = fitz.open(r"{}/{}".format(figurePath, newFigPdf))
            # iterate over PDF pages
            for page_index in range(len(pdf_file)):
                # get the page itself
                page = pdf_file[page_index]
                image_list = page.getImageList()
                # printing number of images found in this page
                if image_list:
                    print(f"[+] Found a total of {len(image_list)} images in page {page_index}")
                else:
                    print("[!] No images found on page", page_index)
                for image_index, img in enumerate(page.getImageList(), start=1):
                    # get the XREF of the image
                    xref = img[0]
                    # extract the image bytes
                    base_image = pdf_file.extractImage(xref)
                    image_bytes = base_image["image"]
                    # get the image extension
                    image_ext = base_image["ext"]
                    # load it to PIL
                    image = Image.open(io.BytesIO(image_bytes))
                    # save it to local disk
                    image.save(open(f"{figurePath}/{newFigPdf1[-2]}.tiff", "wb"))
                    # im = Image.open(f'{filePath}\{newFigPdf1[-2]}.png')
                    # im_mirror = ImageOps.mirror(im)
                    # im_mirror = im_mirror.rotate(180)
                    # im.save(f'{filePath}\{newFigPdf1[-2]}.tiff')
                    pdf_file.close()
                    # print(r"{}".format(figPDF),"UUUUU")
                    # print(r"{}\{}".format(filePath, newFigPdf))
                    os.remove(r"{}/{}".format(figurePath, newFigPdf))
        # 2. change its extension accordingly
        # 3. once image extracted and renamed, delete it

    # Image_renaming and format conversion
    def Image_auto(self, figurePath, isbn):
        allImagesList = []
        allImagesList = os.listdir(r"{}".format(figurePath))
        allImagesList = [figurePath + i for i in allImagesList]
        # acceptable_format = ["a","b","c"]
        # other = [i for i in allImagesList if i in acceptable_format]
        pdfList = [i for i in allImagesList if i.split(".")[-1] == "pdf"]
        otherImgList = [i for i in allImagesList if (i.split(".")[-1] != "pdf" and i.split(".")[-1] != "xlsx" and i.split(".")[-1] != 'pptx')]
        # print(allImagesList)
        # print("pdfList: ",pdfList)
        # print(otherImgList,)

        if len(pdfList) > 0:
            try:
                self.pdfProcess(pdfList, isbn)
            except:
                print(" embedded")
                pass

        # for i in otherImgList:
        self.imgRenaming(otherImgList, isbn)

    # def uiData(self,author):
    # pass

    def createReport(self, figurePath, docPath, oldallImagesList, title, chapter, isbn, author,
                     edition, email, figDict):
        print("**************DICT WHICH I AM PASSING WHILE CALLING createReport Function*****************")
        print("-->",figDict)
        print("******************************************************************************************")
        # k=c2.Image_auto(filePath)
        #print("Inside report creation method")
        #print(figurePath,"figurePathhhhhhhhhhh")
        allImagesList = os.listdir(r"{}".format(figurePath))

        k1 = self.fig_captxn_citatxn(r"{}".format(docPath))
        #print(k1,"jnjvnvdjnv")
        figure = []
        fig_no = []
        citation = []
        allImagesList = [i for i in allImagesList if (i.split(".")[-1] != "xlsx" and i.split(".")[-1] != 'pptx')]
        # for i in allImagesList:
        #     print(figurePath+i,"checkkkkkkkkkk")
        #     print(Image.open(figurePath+i).info['dpi'],"dpiiiiiiiiiiiiiiiiii")
        oldallImagesList = [i for i in oldallImagesList if (i.split(".")[-1] != "xlsx" and i.split(".")[-1] != 'pptx')]
        k = []
        for fig in allImagesList:
            if re.findall('fig[A-Za-z]', fig.lower()):
                os.remove(r"{}/{}".format(figurePath, fig))
                k.append(fig)

           # print(fig, "dddddddddddddddd", figurePath)
        for i in k:
            allImagesList.remove(i)
        k = []
        for fig in oldallImagesList:
            if re.findall('fig[A-Za-z]', fig.lower()):
                # os.remove(r"{}/{}".format(figurePath, fig))
                k.append(fig)

            #print(fig, "dddddddddddddddd", figurePath)
        for i in k:
            oldallImagesList.remove(i)


        allImagesList.sort(key=lambda f: int(re.sub('\D', '', f)))
        for i in allImagesList:
            try:
                #print(figurePath + i, "checkkkkkkkkkk")
                #print(Image.open(figurePath + i).info['dpi'][0], "dpiiiiiiiiiiiiiiiiii")
                dpi=Image.open(figurePath + i).info['dpi'][0]

            except:
                dpi="straight pack away"
                #print("straight pack away")
        oldallImagesList.sort(key=lambda f: int(re.sub('\D', '', f)))

        for a in k1:
            #print(a)
            figure.append(a[0])
            fi=a[0].split(" ")[1]
            #print(fi,re.search("\.",fi),"dddddddddddddddddddddddd")
            match = ""
            if re.findall("\.",fi):
                match="dot"
            elif re.findall("-",fi):
                match="hyphen"
            #print(match)
            try:
                if match=="dot":
                    #print(a[0],"if")
                    fig_no.append(int(a[0].split(" ")[1].split(".")[1]))
                elif match=="hyphen":
                    #print(a,"else")
                    fig_no.append(int(a[0].split(" ")[1].split("-")[1]))
                else:
                    continue
            except Exception as e:
                #print("Error belonging to k1:--> ", e)
                a = list(a)
                a[0] = a[0].replace("\u2003", " ")
                a[0] = a[0].replace("\u2002", " ")

                if match == "dot":
                    fig_no.append(int(a[0].split(" ")[1].split(".")[1]))
                elif match == "hyphen":
                    fig_no.append(int(a[0].split(" ")[1].split("-")[1]))
                else:
                    continue

            citation.append(a[2])
        # print(fig_no, "fig_no ",len(fig_no))
        figNum = []

        for i in allImagesList:
            #print(i,"(((((",i.split("-"),"********",i.split("-")[1],"********************",int(i.split("-")[1]))
            # print(i)
            try:
                try:
                    #int(i.split("-")[1]) == int(i.split("-")[2]):
                    figNum.append(int(i.split("-")[2]))
                    #print("22222222222222222222222")
                except:
                    figNum.append(int(i.split("-")[1]))
                    #print("11111111111111111111111111111111")
            except:
                #print(i, (int(i.split("-")[1])), 'wrong format')
                pass
        #print(figNum, "fig num",len(figNum))
        seen = set()
        uniq = [x for x in figNum if x in seen or seen.add(x)]
        # print(seen, " seen ")
        # print(uniq, "uniq",len(uniq))
        citList = []
        # print(len(citation),"citation ")
        for i in uniq:
            #print(i,"*..")
            # #fig_no.insert(fig_no.index(i), i)
            # fii=fig_no.index(i)
            #
            # ci=citation[fii]
            # print(fii, 'fii', i,'ci',ci)
            # citList.append(ci)
            citList.append(citation[fig_no.index(i)])

        for i in uniq:
            fig_no.insert(fig_no.index(i), i)

        for i in citList:
            citation.insert(citation.index(i), i)
        #print(len(citation), "citation1 ",len(figNum),'*0  ',len(fig_no))
        not_common = (set(fig_no) - (set(figNum)))
        # print(not_common,"not_common")

        for i in not_common:
            if i in fig_no:

                allImagesList.insert(fig_no.index(i), '')
                oldallImagesList.insert(fig_no.index(i), '')
            else:
                figure.insert(figNum.index(i), '')
                citation.insert(figNum.index(i), '')
        dpi=[]                                                                  # change
        for j in allImagesList:		# change
            #print(j.split(".")[-1],"dpiii")		# change
            try:# j.split(".")[-1]=="eps" and j.split(".")[-1]=="pdf":	# change
                #print(figurePath + j, "checkkkkkkkkkk")	# change
                #print(Image.open(figurePath + j).info['dpi'][0], "dpiiiiiiiiiiiiiiiiii")	# change
                dpi.append(int(Image.open(figurePath + j).info['dpi'][0]))	# change
            except:			# change

                dpi.append("straight pack away")		# change
                #print("straight pack away")		# change
        #print(dpi,"hhh")                                                         # change
        x = 0
        filePathxl = r"" + settings.MEDIA_ROOT + "/created_report/HS_TA_Template.xlsx"
        wb = xl.load_workbook(filePathxl)
        sheet = wb.active
        for i in range(1, 56):
            if sheet.cell(i, 1).value == "TOTALS":
                x = i
                break

        for i in range(len(citation)):
            x = x + 1
            sheet.cell(x, 1).value = fig_no[i]
            sheet.cell(x, 2).value = 1
            figDict1 = {}
        k = figDict.keys()
        # print(k)
        for i in k:
            try:
                if re.findall("\.", i):
                    figDict1[i.split(".")[1]] = figDict[i]
                elif re.findall("-", i):
                    figDict1[i.split("-")[1]] = figDict[i]
                else:
                    figDict1[i] = figDict[i]
            except:
                pass
        #print(figDict1)
        for i in range(len(citation)):
            x = x + 1
            sheet.cell(x, 1).value = fig_no[i]

            sheet.cell(x, 2).value = 1
            try:
                if len(allImagesList[i].split('-')) == 4:
                    sheet.cell(x, 11).value = "EMSS Note: "+ figDict1[str(fig_no[i])+''+allImagesList[i].split('-')[2]]
                    #print(str(fig_no[i])+''+allImagesList[i].split('-')[2]," hhhhhhhh ")
                else:
                    sheet.cell(x, 11).value = "EMSS Note: "+ figDict1[str(fig_no[i])]
            except:
                pass



            # sheet.cell(x, 5).value = citation[i]
            #print(citation[i])
            if citation[i] != "":  # base condition
                if re.search("^From", citation[i]):
                    sheet.cell(x, 5).value = citation[i]
                else:
                    if  re.search("\(From", citation[i]):
                        ct = citation[i].rsplit("(From", 1)
                    elif re.search("From", citation[i]):
                        ct = citation[i].rsplit("From", 1)

                    sheet.cell(x, 5).value = "From" + ct[1]

            else:
                sheet.cell(x, 5).value = citation[i]

            print("--------------------IMPORTANT LIST-------------------------------------------")
            print(allImagesList)
            print("-----------------------------------------------------------------------------")
            print(oldallImagesList)
            print("=============================================================================")
            if allImagesList[i] == "" and fig_no[i] < 10:
                allImagesList[i] = "f" + "chapter" + "-" + "0" + str(fig_no[i]) + "-" + isbn + ".eps"
            elif allImagesList[i] == "" and fig_no[i] >= 10:
                allImagesList[i] = "f" + "chapter" + "-" + str(fig_no[i]) + "-" + isbn + ".eps"

            sheet.cell(x, 7).value = allImagesList[i]

            sheet.cell(x, 6).value = oldallImagesList[i]

            sheet.cell(x, 13).value = dpi[i]   # change

        today = datetime.date.today()
        date = "{:%d-%b-%Y}".format(today)
        sheet.cell(2, 2).value = author
        sheet.cell(3, 2).value = title
        sheet.cell(4, 9).value = chapter
        sheet.cell(2, 9).value = isbn
        sheet.cell(5, 9).value = date
        sheet.cell(5, 2).value = email
        sheet.cell(4, 2).value = "Aptara"
        sheet.cell(3, 9).value = edition

        wb.save(filePathxl)

            # try:
            #     sheet.cell(x, 13).value = int(j)
            # except:
            #     sheet.cell(x, 13).value = j



        wb.save(filePathxl)

        # for i in range(len(citation)):
        #     print(fig_no[i], " ", allImagesList[i],' ',oldallImagesList[i],' ', citation[i])